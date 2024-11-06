from docx import Document


class WordFileReader:
    @staticmethod
    def read_file(file_path: str) -> list[str]:
        doc = Document(file_path)

        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        return content

    @staticmethod
    def write_file(file_path: str, new_text: list[str]) -> None:
        doc = Document(file_path)

        for paragraph in doc.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

        for line in new_text:
            doc.add_paragraph(line)

        doc.save(file_path)
